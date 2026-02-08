# biography_publisher.py - WITH CLEAN CONFETTI AND EXPORT OPTIONS
import streamlit as st
import json
import base64
from datetime import datetime
import time
from io import BytesIO

# ============================================================================
# DOCX LIBRARY IMPORT
# ============================================================================
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Page setup
st.set_page_config(page_title="Biography Publisher", layout="wide")
st.title("📖 Beautiful Biography Creator")

# Custom CSS for better styling
st.markdown("""
<style>
    .main-title {
        text-align: center;
        color: #2c3e50;
        font-size: 3em;
        margin-bottom: 0.5em;
    }
    .subtitle {
        text-align: center;
        color: #7f8c8d;
        font-size: 1.2em;
        margin-bottom: 2em;
    }
    .story-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border-left: 4px solid #3498db;
    }
    .download-btn {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 1rem 2rem;
        border-radius: 8px;
        font-size: 1.1em;
        font-weight: 600;
        cursor: pointer;
        width: 100%;
        margin-top: 1rem;
        transition: transform 0.2s;
    }
    .download-btn:hover {
        transform: translateY(-2px);
    }
    .stats-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        text-align: center;
        border-top: 4px solid #2ecc71;
    }
    .preview-box {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px dashed #dee2e6;
        font-family: 'Courier New', monospace;
        font-size: 0.9em;
        line-height: 1.6;
    }
    .format-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin-bottom: 1rem;
        border-top: 4px solid #3498db;
    }
    .format-icon {
        font-size: 2.5em;
        margin-bottom: 0.5rem;
    }
    .export-option {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px solid #e9ecef;
        margin-bottom: 1.5rem;
    }
</style>
""", unsafe_allow_html=True)

def show_celebration():
    """Show celebration effects - FIXED CONFETTI CODE"""
    # Balloons
    st.balloons()
    
    # FIXED Confetti effect - no trailing JavaScript visible
    st.markdown("""
    <style>
    @keyframes confetti-fall {
        0% { transform: translateY(-100px) rotate(0deg); opacity: 1; }
        100% { transform: translateY(100vh) rotate(360deg); opacity: 0; }
    }
    .confetti {
        position: fixed;
        width: 10px;
        height: 10px;
        background: #ff6b6b;
        animation: confetti-fall 3s linear forwards;
        z-index: 1000;
    }
    </style>
    <div id="confetti-container"></div>
    <script>
    (function createConfetti() {
        const container = document.getElementById('confetti-container');
        const colors = ['#ff6b6b', '#4ecdc4', '#45b7d1', '#96ceb4', '#ffeaa7'];
        
        for (let i = 0; i < 50; i++) {
            const confetti = document.createElement('div');
            confetti.className = 'confetti';
            confetti.style.left = Math.random() * 100 + 'vw';
            confetti.style.animationDelay = Math.random() * 2 + 's';
            confetti.style.backgroundColor = colors[Math.floor(Math.random() * colors.length)];
            confetti.style.width = Math.random() * 10 + 5 + 'px';
            confetti.style.height = Math.random() * 10 + 5 + 'px';
            container.appendChild(confetti);
            
            // Remove element after animation completes
            setTimeout(() => {
                if (confetti.parentNode === container) {
                    container.removeChild(confetti);
                }
            }, 5000);
        }
        
        // Remove container after all confetti is gone
        setTimeout(() => {
            if (container.parentNode) {
                container.parentNode.removeChild(container);
            }
        }, 6000);
    })();
    </script>
    """, unsafe_allow_html=True)
    
    # Success message with animation
    st.markdown("""
    <div style="text-align: center; padding: 2rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 15px; color: white; margin: 2rem 0;">
        <h1 style="color: white; font-size: 2.5em;">🎉 Biography Created! 🎉</h1>
        <p style="font-size: 1.2em;">Your beautiful life story is ready to share!</p>
    </div>
    """, unsafe_allow_html=True)

def decode_stories_from_url():
    """Extract stories from URL parameter - FIXED FOR COMPATIBILITY"""
    try:
        # Try new method first (Streamlit 1.28+)
        if hasattr(st, 'query_params'):
            query_params = st.query_params.to_dict()
            encoded_data = query_params.get("data")
            
            if isinstance(encoded_data, list):
                encoded_data = encoded_data[0]
        else:
            # Fall back to experimental method
            query_params = st.experimental_get_query_params()
            encoded_data = query_params.get("data", [None])[0]
        
        if not encoded_data:
            return None
            
        # Decode the data
        json_data = base64.b64decode(encoded_data).decode()
        stories_data = json.loads(json_data)
        
        return stories_data
    except Exception as e:
        st.error(f"Error loading data: {str(e)}")
        return None

# ============================================================================
# NEW: EXPORT OPTION FUNCTIONS
# ============================================================================
def create_docx_biography(stories_data, include_questions=True):
    """Create a professionally formatted Word document (.docx) with option for questions"""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available. Please install with: pip install python-docx==1.1.0")
    
    # Extract data
    user_name = stories_data.get("user", "Unknown")
    user_profile = stories_data.get("user_profile", {})
    stories_dict = stories_data.get("stories", {})
    
    # Get author name
    if user_profile and 'first_name' in user_profile:
        first_name = user_profile.get('first_name', '')
        last_name = user_profile.get('last_name', '')
        author_name = f"{first_name} {last_name}".strip()
        if not author_name:
            author_name = user_name
    else:
        author_name = user_name
    
    # Create document
    doc = Document()
    
    # ========== SET UP DOCUMENT STYLES ==========
    
    # Title style - Check if exists first
    try:
        title_style = doc.styles['CustomTitle']
    except KeyError:
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri Light'
        title_font.size = Pt(28)
        title_font.bold = True
        title_font.color.rgb = RGBColor(44, 82, 130)  # Dark blue
    
    # Heading 1 style
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Calibri'
    heading1_style.font.size = Pt(20)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(44, 82, 130)
    
    # Heading 2 style
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = 'Calibri'
    heading2_style.font.size = Pt(16)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(66, 133, 244)  # Blue
    
    # Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    # Quote style - Only create if it doesn't exist (FIXED)
    try:
        quote_style = doc.styles['Quote']
    except KeyError:
        quote_style = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
        quote_style.font.name = 'Calibri'
        quote_style.font.size = Pt(11)
        quote_style.font.italic = True
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.right_indent = Inches(0.5)
    
    # ========== CREATE COVER PAGE ==========
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("TELL MY STORY\n")
    title_run.font.name = 'Calibri Light'
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(44, 82, 130)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run("A Personal Biography\n")
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(20)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n\n\n\n")
    
    # Author name
    author_para = doc.add_paragraph()
    author_run = author_para.add_run(f"The Life Story of\n{author_name.upper()}")
    author_run.font.name = 'Calibri'
    author_run.font.size = Pt(24)
    author_run.font.color.rgb = RGBColor(0, 0, 0)
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n\n\n\n\n\n")
    
    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Compiled on {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    date_run.italic = True
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Page break
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    
    toc_title = doc.add_heading('TABLE OF CONTENTS', 1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Collect all sessions for TOC
    try:
        sorted_sessions = sorted(stories_dict.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 0)
    except:
        sorted_sessions = stories_dict.items()
    
    for session_id, session_data in sorted_sessions:
        session_title = session_data.get("title", f"Chapter {session_id}")
        questions = session_data.get("questions", {})
        
        if questions:
            para = doc.add_paragraph()
            para.style = 'Normal'
            run = para.add_run(f"{session_title}")
            run.bold = True
            run.font.size = Pt(12)
            
            # Add page numbers placeholder
            para.add_run(f"\t\t\t...... ")
    
    doc.add_paragraph("\n")
    
    # ========== INTRODUCTION ==========
    
    intro_title = doc.add_heading('INTRODUCTION', 1)
    intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    intro_para = doc.add_paragraph()
    export_type = "Interview Q&A" if include_questions else "Biography"
    intro_text = f"This {export_type.lower()} captures the unique life journey of {author_name}, "
    intro_text += f"compiled from personal reflections shared on {datetime.now().strftime('%B %d, %Y')}. "
    intro_text += "Each chapter represents a different phase of life, preserved here for future generations."
    intro_para.add_run(intro_text)
    
    doc.add_page_break()
    
    # ========== CHAPTERS AND STORIES ==========
    
    chapter_num = 0
    total_stories = 0
    total_words = 0
    
    for session_id, session_data in sorted_sessions:
        session_title = session_data.get("title", f"Chapter {session_id}")
        questions = session_data.get("questions", {})
        
        if not questions:
            continue
        
        chapter_num += 1
        
        # Chapter header
        chapter_title = doc.add_heading(f'CHAPTER {chapter_num}: {session_title.upper()}', 1)
        chapter_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Process each story in this chapter
        story_num = 0
        for question, answer_data in questions.items():
            if isinstance(answer_data, dict):
                answer = answer_data.get("answer", "")
                date_recorded = answer_data.get("timestamp", datetime.now().isoformat())[:10]
            else:
                answer = str(answer_data)
                date_recorded = datetime.now().isoformat()[:10]
            
            if not answer.strip():
                continue
            
            story_num += 1
            total_stories += 1
            word_count = len(answer.split())
            total_words += word_count
            
            # Story header - only include question if option is selected
            if include_questions:
                story_header = doc.add_heading(f'Story {story_num}: {question}', 2)
            else:
                story_header = doc.add_heading(f'Story {story_num}', 2)
            
            # Date if available
            if date_recorded:
                date_para = doc.add_paragraph()
                date_run = date_para.add_run(f"Recorded: {date_recorded}")
                date_run.font.size = Pt(10)
                date_run.font.color.rgb = RGBColor(100, 100, 100)
                date_run.italic = True
            
            # Story content
            content_para = doc.add_paragraph()
            content_para.add_run(answer.strip())
            
            # Word count
            count_para = doc.add_paragraph()
            count_run = count_para.add_run(f"[{word_count} words]")
            count_run.font.size = Pt(9)
            count_run.font.color.rgb = RGBColor(150, 150, 150)
            
            doc.add_paragraph()  # Add spacing between stories
        
        doc.add_page_break()  # New page for next chapter
    
    # ========== STATISTICS PAGE ==========
    
    stats_title = doc.add_heading('BIOGRAPHY STATISTICS', 1)
    stats_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Create a table for stats
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    # Data rows
    export_type_display = "Interview Q&A" if include_questions else "Biography"
    metrics = [
        ('Export Type', export_type_display),
        ('Total Chapters', str(chapter_num)),
        ('Total Stories', str(total_stories)),
        ('Total Words', f"{total_words:,}"),
        ('Average Story Length', f"{total_words//total_stories if total_stories > 0 else 0} words"),
        ('Compiled Date', datetime.now().strftime('%B %d, %Y')),
        ('Compiled Time', datetime.now().strftime('%I:%M %p'))
    ]
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    doc.add_paragraph("\n\n")
    
    # Conclusion
    conclusion_para = doc.add_paragraph()
    conclusion_text = f"This {export_type_display.lower()} contains {total_stories} personal stories from {author_name}'s life, "
    conclusion_text += f"totaling {total_words:,} words across {chapter_num} chapters. "
    conclusion_text += "These memories are now preserved for future generations to cherish."
    conclusion_para.add_run(conclusion_text)
    
    doc.add_paragraph("\n")
    
    # Footer note
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run("Created with Tell My Story Biographer")
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========== SAVE TO BYTESIO ==========
    
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes, author_name, chapter_num, total_stories, total_words

def create_beautiful_biography(stories_data, include_questions=True):
    """Create a professionally formatted biography with option for questions"""
    user_name = stories_data.get("user", "Unknown")
    user_profile = stories_data.get("user_profile", {})
    stories_dict = stories_data.get("stories", {})
    summary = stories_data.get("summary", {})
    
    # Get display name
    if user_profile and 'first_name' in user_profile:
        first_name = user_profile.get('first_name', '')
        last_name = user_profile.get('last_name', '')
        display_name = f"{first_name} {last_name}".strip()
        if not display_name:
            display_name = user_name
    else:
        display_name = user_name
    
    # Collect all stories
    all_stories = []
    try:
        # Sort sessions numerically
        sorted_sessions = sorted(stories_dict.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 0)
    except:
        sorted_sessions = stories_dict.items()
    
    for session_id, session_data in sorted_sessions:
        session_title = session_data.get("title", f"Chapter {session_id}")
        
        for question, answer_data in session_data.get("questions", {}).items():
            if isinstance(answer_data, dict):
                answer = answer_data.get("answer", "")
            else:
                answer = str(answer_data)
                
            if answer.strip():
                all_stories.append({
                    "session": session_title,
                    "question": question,
                    "answer": answer,
                    "date": answer_data.get("timestamp", datetime.now().isoformat())[:10] 
                             if isinstance(answer_data, dict) else datetime.now().isoformat()[:10],
                    "session_id": session_id
                })
    
    if not all_stories:
        return "No stories found to publish.", [], display_name, 0, 0, 0
    
    # ========== CREATE BEAUTIFUL BIOGRAPHY ==========
    export_type = "INTERVIEW Q&A" if include_questions else "BIOGRAPHY"
    bio_text = "=" * 70 + "\n"
    bio_text += f"{'TELL MY STORY':^70}\n"
    bio_text += f"{export_type:^70}\n"
    bio_text += "=" * 70 + "\n\n"
    
    bio_text += f"THE LIFE STORY OF\n{display_name.upper()}\n\n"
    bio_text += "-" * 70 + "\n\n"
    
    # Personal Information
    if user_profile:
        bio_text += "PERSONAL INFORMATION\n"
        bio_text += "-" * 40 + "\n"
        if user_profile.get('birthdate'):
            bio_text += f"Date of Birth: {user_profile.get('birthdate')}\n"
        if user_profile.get('gender'):
            bio_text += f"Gender: {user_profile.get('gender')}\n"
        bio_text += "\n"
    
    # Table of Contents
    bio_text += "TABLE OF CONTENTS\n"
    bio_text += "-" * 40 + "\n\n"
    
    current_session = None
    chapter_num = 0
    for story in all_stories:
        if story["session"] != current_session:
            chapter_num += 1
            bio_text += f"Chapter {chapter_num}: {story['session']}\n"
            current_session = story["session"]
    
    bio_text += "\n" + "=" * 70 + "\n\n"
    
    # Introduction
    bio_text += "INTRODUCTION\n\n"
    bio_text += f"This {export_type.lower()} captures the unique life journey of {display_name}, "
    bio_text += f"compiled from personal reflections shared on {datetime.now().strftime('%B %d, %Y')}. "
    bio_text += "Each chapter represents a different phase of life, preserved here for future generations.\n\n"
    
    bio_text += "=" * 70 + "\n\n"
    
    # Chapters with stories
    current_session = None
    chapter_num = 0
    story_num = 0
    
    for story in all_stories:
        if story["session"] != current_session:
            chapter_num += 1
            bio_text += "\n" + "=" * 70 + "\n"
            bio_text += f"CHAPTER {chapter_num}: {story['session'].upper()}\n"
            bio_text += "=" * 70 + "\n\n"
            current_session = story["session"]
        
        story_num += 1
        
        # Include question only if selected
        if include_questions:
            bio_text += f"Story {story_num}\n"
            bio_text += f"Topic: {story['question']}\n"
        else:
            bio_text += f"Story {story_num}\n"
        
        if story['date']:
            bio_text += f"Recorded: {story['date']}\n"
        
        bio_text += "-" * 40 + "\n"
        
        # Format the answer with proper paragraphs
        answer = story['answer'].strip()
        paragraphs = answer.split('\n')
        
        for para in paragraphs:
            if para.strip():
                bio_text += f"{para.strip()}\n\n"
        
        bio_text += "\n"
    
    # Conclusion
    bio_text += "=" * 70 + "\n\n"
    bio_text += "CONCLUSION\n\n"
    bio_text += f"This collection contains {story_num} stories across {chapter_num} chapters, "
    bio_text += f"each one a piece of {display_name}'s unique mosaic of memories. "
    bio_text += "These reflections will continue to resonate long into the future.\n\n"
    
    # Statistics
    bio_text += "-" * 70 + "\n"
    bio_text += f"{export_type} STATISTICS\n"
    bio_text += "-" * 40 + "\n"
    bio_text += f"• Export Type: {export_type}\n"
    bio_text += f"• Total Stories: {story_num}\n"
    bio_text += f"• Total Chapters: {chapter_num}\n"
    
    # Calculate word count
    total_words = sum(len(story['answer'].split()) for story in all_stories)
    bio_text += f"• Total Words: {total_words:,}\n"
    
    # Find longest story
    if all_stories:
        longest = max(all_stories, key=lambda x: len(x['answer'].split()))
        bio_text += f"• Longest Story: \"{longest['question'][:50]}...\" ({len(longest['answer'].split())} words)\n"
    
    bio_text += f"• Compiled: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n"
    bio_text += "-" * 70 + "\n\n"
    
    # Final note
    bio_text += "This digital legacy was created with Tell My Story Biographer.\n\n"
    bio_text += "=" * 70
    
    return bio_text, all_stories, display_name, story_num, chapter_num, total_words

def create_html_biography(stories_data, include_questions=True):
    """Create an HTML version with option for questions"""
    bio_text, all_stories, display_name, story_num, chapter_num, total_words = create_beautiful_biography(stories_data, include_questions)
    
    export_type = "Interview Q&A" if include_questions else "Biography"
    
    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{display_name}'s {export_type}</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Crimson+Text:ital,wght@0,400;0,600;0,700;1,400&family=Open+Sans:wght@300;400;600&display=swap');
        
        body {{
            font-family: 'Crimson Text', serif;
            line-height: 1.8;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            background: #fefefe;
        }}
        .header {{
            text-align: center;
            padding: 40px 0;
            border-bottom: 3px double #2c5282;
            margin-bottom: 40px;
        }}
        h1 {{
            font-size: 2.8em;
            color: #2c5282;
            margin-bottom: 10px;
        }}
        .subtitle {{
            font-family: 'Open Sans', sans-serif;
            font-size: 1.2em;
            color: #666;
        }}
        .chapter {{
            margin: 50px 0;
        }}
        .chapter-title {{
            color: #2c5282;
            border-bottom: 2px solid #e2e8f0;
            padding-bottom: 10px;
            margin-bottom: 30px;
        }}
        .story {{
            margin: 30px 0;
            padding: 25px;
            background: #f8fafc;
            border-radius: 8px;
            border-left: 4px solid #4299e1;
        }}
        .question {{
            font-weight: 700;
            color: #2d3748;
            margin-bottom: 15px;
            font-size: 1.2em;
        }}
        .answer {{
            white-space: pre-line;
            font-size: 1.1em;
        }}
        .stats {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            border-radius: 15px;
            text-align: center;
            margin: 40px 0;
        }}
        .stat-item {{
            display: inline-block;
            margin: 0 20px;
        }}
        .stat-number {{
            font-size: 2.5em;
            font-weight: 700;
            display: block;
        }}
        .footer {{
            text-align: center;
            margin-top: 50px;
            padding-top: 20px;
            border-top: 2px solid #e2e8f0;
            color: #718096;
        }}
        @media print {{
            body {{ padding: 0; }}
            .no-print {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{display_name}'s Life Story</h1>
        <div class="subtitle">{export_type} • {datetime.now().strftime('%B %d, %Y')}</div>
    </div>
    
    <div class="stats">
        <div class="stat-item">
            <span class="stat-number">{chapter_num}</span>
            <span>Chapters</span>
        </div>
        <div class="stat-item">
            <span class="stat-number">{story_num}</span>
            <span>Stories</span>
        </div>
        <div class="stat-item">
            <span class="stat-number">{total_words:,}</span>
            <span>Words</span>
        </div>
    </div>
    
    <div class="content">
'''

    # Add chapters
    current_session = None
    chapter_num = 0
    
    for story in all_stories:
        if story["session"] != current_session:
            chapter_num += 1
            current_session = story["session"]
            html += f'''
            <div class="chapter">
                <h2 class="chapter-title">Chapter {chapter_num}: {story['session']}</h2>
            '''
        
        html += f'''
        <div class="story">
        '''
        
        # Include question only if selected
        if include_questions:
            html += f'''
            <div class="question">✏️ {story['question']}</div>
            '''
        
        html += f'''
            <div class="answer">{story['answer']}</div>
        '''
        
        if story['date']:
            html += f'''
            <div style="margin-top: 15px; font-size: 0.9em; color: #718096;">
                Recorded: {story['date']}
            </div>
            '''
        
        html += '</div>'
        
        # Close chapter if next story is different session or last story
        next_idx = all_stories.index(story) + 1
        if next_idx >= len(all_stories) or all_stories[next_idx]["session"] != current_session:
            html += '</div>'

    html += f'''
    </div>
    
    <div class="footer">
        <p>Created with Tell My Story Biographer • {export_type}</p>
        <p>{datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
    </div>
    
    <div class="no-print" style="text-align: center; margin-top: 40px;">
        <button onclick="window.print()" style="
            background: #48bb78;
            color: white;
            border: none;
            padding: 12px 30px;
            border-radius: 8px;
            font-size: 1.1em;
            cursor: pointer;
            margin: 20px;
        ">
            🖨️ Print This {export_type}
        </button>
    </div>
</body>
</html>'''
    
    return html, display_name

# ============================================================================
# MAIN APP INTERFACE
# ============================================================================

st.markdown('<h1 class="main-title">📖 Beautiful Biography Creator</h1>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">Transform your life stories into a professionally formatted book</p>', unsafe_allow_html=True)

# Format information card
st.markdown("""
<div class="format-card">
    <h3>🎯 Available Export Formats</h3>
    <div style="display: grid; grid-template-columns: repeat(4, 1fr); gap: 20px; margin-top: 1.5rem;">
        <div style="text-align: center; padding: 1rem; background: #f8f9fa; border-radius: 10px;">
            <div class="format-icon">📄</div>
            <div><strong>TEXT</strong></div>
            <div style="font-size: 0.9em; color: #666;">Plain Text File</div>
            <div style="font-size: 0.8em; color: #888; margin-top: 0.5rem;">.txt</div>
        </div>
        <div style="text-align: center; padding: 1rem; background: #f8f9fa; border-radius: 10px;">
            <div class="format-icon">🌐</div>
            <div><strong>HTML</strong></div>
            <div style="font-size: 0.9em; color: #666;">Web Format</div>
            <div style="font-size: 0.8em; color: #888; margin-top: 0.5rem;">.html</div>
        </div>
        <div style="text-align: center; padding: 1rem; background: #f8f9fa; border-radius: 10px;">
            <div class="format-icon">📝</div>
            <div><strong>MARKDOWN</strong></div>
            <div style="font-size: 0.9em; color: #666;">Easy Editing</div>
            <div style="font-size: 0.8em; color: #888; margin-top: 0.5rem;">.md</div>
        </div>
        <div style="text-align: center; padding: 1rem; background: #e8f4f8; border-radius: 10px; border: 2px solid #3498db;">
            <div class="format-icon">📘</div>
            <div><strong>WORD DOC</strong></div>
            <div style="font-size: 0.9em; color: #666;">Microsoft Word</div>
            <div style="font-size: 0.8em; color: #3498db; margin-top: 0.5rem; font-weight: bold;">.docx</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# Try to get data from URL first
stories_data = decode_stories_from_url()

if stories_data:
    # Auto-process data from URL
    user_name = stories_data.get("user", "Unknown")
    user_profile = stories_data.get("user_profile", {})
    summary = stories_data.get("summary", {})
    
    # Count stories
    story_count = 0
    for session_id, session_data in stories_data.get("stories", {}).items():
        story_count += len(session_data.get("questions", {}))
    
    if story_count > 0:
        # Display user info
        col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
        
        with col1:
            if user_profile and user_profile.get('first_name'):
                st.success(f"✅ Welcome, **{user_profile.get('first_name')} {user_profile.get('last_name', '')}**!")
            else:
                st.success(f"✅ Welcome, **{user_name}**!")
            
            if user_profile and user_profile.get('birthdate'):
                st.caption(f"🎂 Born: {user_profile.get('birthdate')}")
        
        with col2:
            st.metric("Sessions", len(stories_data.get("stories", {})))
        
        with col3:
            st.metric("Stories", story_count)
        
        with col4:
            # Count total words
            total_words = 0
            for session_id, session_data in stories_data.get("stories", {}).items():
                for question, answer_data in session_data.get("questions", {}).items():
                    if isinstance(answer_data, dict):
                        answer = answer_data.get("answer", "")
                    else:
                        answer = str(answer_data)
                    total_words += len(answer.split())
            st.metric("Words", f"{total_words:,}")
        
        # DOCX availability indicator
        if not DOCX_AVAILABLE:
            st.warning("⚠️ DOCX export requires python-docx. Other formats are available.")
        
        # NEW: Export Options Section
        st.markdown("---")
        st.subheader("📝 Export Options")
        
        col_option1, col_option2 = st.columns(2)
        
        with col_option1:
            st.markdown("""
            <div class="export-option">
                <h4>📚 Biography Format</h4>
                <p><strong>Just the answers</strong> - Clean narrative format</p>
                <p>• Flowing story without questions</p>
                <p>• More like a traditional biography</p>
                <p>• Professional for sharing</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col_option2:
            st.markdown("""
            <div class="export-option">
                <h4>🎤 Interview Format</h4>
                <p><strong>Questions & answers</strong> - Complete Q&A format</p>
                <p>• Includes all interview questions</p>
                <p>• Preserves the conversational flow</p>
                <p>• Great for reference</p>
            </div>
            """, unsafe_allow_html=True)
        
        # Radio button for export option
        export_format = st.radio(
            "Choose your export format:",
            ["📚 Biography Format (Just Answers)", "🎤 Interview Format (Questions & Answers)"],
            index=0,
            horizontal=True,
            key="export_format"
        )
        
        include_questions = export_format == "🎤 Interview Format (Questions & Answers)"
        
        # Generate biography button
        if st.button("✨ Create Beautiful Biography", type="primary", use_container_width=True, key="create_bio_btn"):
            with st.spinner("🖋️ Crafting your beautiful biography..."):
                time.sleep(1)
                
                # Create all versions
                bio_text, all_stories, author_name, story_num, chapter_num, total_words = create_beautiful_biography(
                    stories_data, include_questions
                )
                html_bio, html_name = create_html_biography(stories_data, include_questions)
                
                # Create DOCX if available
                docx_data = None
                if DOCX_AVAILABLE:
                    try:
                        docx_bytes, docx_name, docx_chapters, docx_stories, docx_words = create_docx_biography(
                            stories_data, include_questions
                        )
                        docx_data = docx_bytes
                    except Exception as e:
                        st.error(f"⚠️ DOCX creation failed: {str(e)}")
                        docx_data = None
            
            # Show celebration
            show_celebration()
            
            # Show preview
            export_type_display = "Interview Q&A" if include_questions else "Biography"
            st.subheader(f"📖 Your {export_type_display} Preview")
            
            col_preview1, col_preview2 = st.columns([2, 1])
            
            with col_preview1:
                st.markdown('<div class="preview-box">', unsafe_allow_html=True)
                preview_text = bio_text[:1500] + "..." if len(bio_text) > 1500 else bio_text
                st.text(preview_text)
                st.markdown('</div>', unsafe_allow_html=True)
                st.caption(f"Preview of {len(preview_text):,} characters out of {len(bio_text):,} total")
            
            with col_preview2:
                st.markdown('<div class="stats-card">', unsafe_allow_html=True)
                st.metric("Format", export_type_display)
                st.metric("📚 Chapters", chapter_num)
                st.metric("📝 Stories", story_num)
                st.metric("📊 Total Words", f"{total_words:,}")
                st.markdown('</div>', unsafe_allow_html=True)
            
            # Download options
            st.subheader("📥 Download Your Biography")
            
            # Create markdown version
            md_bio = bio_text.replace("=" * 70, "#" * 3)
            
            safe_name = author_name.replace(" ", "_")
            file_suffix = "_Interview" if include_questions else "_Biography"
            
            # Row 1: Download buttons
            col_dl1, col_dl2, col_dl3, col_dl4 = st.columns(4)
            
            with col_dl1:
                st.download_button(
                    label="📄 TEXT",
                    data=bio_text,
                    file_name=f"{safe_name}{file_suffix}.txt",
                    mime="text/plain",
                    use_container_width=True,
                    type="primary",
                    help="Plain text format - universal compatibility"
                )
            
            with col_dl2:
                st.download_button(
                    label="🌐 HTML",
                    data=html_bio,
                    file_name=f"{safe_name}{file_suffix}.html",
                    mime="text/html",
                    use_container_width=True,
                    type="secondary",
                    help="Beautiful web format - ready to print"
                )
            
            with col_dl3:
                st.download_button(
                    label="📝 MARKDOWN",
                    data=md_bio,
                    file_name=f"{safe_name}{file_suffix}.md",
                    mime="text/markdown",
                    use_container_width=True,
                    type="secondary",
                    help="Markdown format for easy editing"
                )
            
            with col_dl4:
                if docx_data:
                    st.download_button(
                        label="📘 WORD DOC",
                        data=docx_data,
                        file_name=f"{safe_name}{file_suffix}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="primary",
                        help="Microsoft Word document - professional formatting"
                    )
                else:
                    st.button(
                        "📘 WORD DOC",
                        disabled=True,
                        use_container_width=True,
                        help="Install python-docx to enable Word export"
                    )
            
            # Row 2: Format descriptions
            col_desc1, col_desc2, col_desc3, col_desc4 = st.columns(4)
            
            with col_desc1:
                st.caption("**TEXT** - Universal format")
            
            with col_desc2:
                st.caption("**HTML** - Web & print ready")
            
            with col_desc3:
                st.caption("**MARKDOWN** - Easy editing")
            
            with col_desc4:
                if docx_data:
                    st.caption("**WORD DOC** - Professional")
                else:
                    st.caption("**WORD DOC** - Not available")
            
            # Story preview
            with st.expander("📋 Preview Your Stories", expanded=False):
                try:
                    sorted_sessions = sorted(stories_data.get("stories", {}).items(), 
                                           key=lambda x: int(x[0]) if x[0].isdigit() else 0)
                except:
                    sorted_sessions = stories_data.get("stories", {}).items()
                
                for session_id, session_data in sorted_sessions[:3]:
                    session_title = session_data.get("title", f"Session {session_id}")
                    st.markdown(f"### {session_title}")
                    
                    for question, answer_data in list(session_data.get("questions", {}).items())[:2]:
                        if isinstance(answer_data, dict):
                            answer = answer_data.get("answer", "")
                        else:
                            answer = str(answer_data)
                        
                        if answer.strip():
                            if include_questions:
                                st.markdown(f"**{question}**")
                            else:
                                st.markdown(f"**Story**")
                            st.write(answer[:200] + "..." if len(answer) > 200 else answer)
                            st.caption(f"{len(answer.split())} words")
                            st.divider()
            
            st.success(f"✨ {export_type_display} created! **{story_num} stories** across **{chapter_num} chapters** ({total_words:,} words)")
            
            # Shareable achievement
            st.markdown("---")
            st.markdown("### 🏆 Your Achievement")
            st.markdown(f"""
            <div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); padding: 2rem; border-radius: 15px; color: white; text-align: center; margin: 1rem 0;">
                <h2 style="color: white; margin-bottom: 1rem;">🎉 {export_type_display} Master! 🎉</h2>
                <div style="display: flex; justify-content: center; gap: 3rem; margin: 1.5rem 0;">
                    <div>
                        <div style="font-size: 2.5em; font-weight: bold;">{chapter_num}</div>
                        <div>Chapters</div>
                    </div>
                    <div>
                        <div style="font-size: 2.5em; font-weight: bold;">{story_num}</div>
                        <div>Stories</div>
                    </div>
                    <div>
                        <div style="font-size: 2.5em; font-weight: bold;">{total_words:,}</div>
                        <div>Words</div>
                    </div>
                </div>
                <p style="font-size: 1.1em; margin-top: 1rem;">Your life story is now preserved in {4 if docx_data else 3} formats!</p>
            </div>
            """, unsafe_allow_html=True)
    
    else:
        st.warning(f"Found your profile but no stories yet.")
        st.info("Go back to the main app and save some stories first!")
        
else:
    # Manual upload option
    st.info("📋 **How to create your biography:**")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        ### 🚀 **Automatic Method**
        1. Go to your Tell My Story app
        2. Answer questions and save your responses
        3. Click the **Publish Biography** button
        4. Your stories will automatically appear here
        
        *No file upload needed!*
        """)
        
        st.markdown("""
        <a href="#" target="_blank">
        <button style="
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            padding: 1rem 2rem;
            border-radius: 8px;
            font-size: 1.1em;
            font-weight: 600;
            cursor: pointer;
            width: 100%;
            margin-top: 1rem;
            text-align: center;
            display: block;
        ">
        📖 Go to Tell My Story App
        </button>
        </a>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        ### 📤 **Manual Upload**
        If you have exported your stories as JSON:
        1. Download the JSON file from the main app
        2. Upload it here:
        """)
        
        uploaded_file = st.file_uploader("Choose a JSON file", type=['json'], label_visibility="collapsed")
        if uploaded_file:
            try:
                uploaded_data = json.load(uploaded_file)
                story_count = sum(len(session.get("questions", {})) for session in uploaded_data.get("stories", {}).values())
                st.success(f"✅ Loaded {story_count} stories")
                
                # Add export format option for manual upload
                manual_export_format = st.radio(
                    "Choose export format for manual upload:",
                    ["📚 Biography Format (Just Answers)", "🎤 Interview Format (Questions & Answers)"],
                    index=0,
                    horizontal=True,
                    key="manual_export_format"
                )
                manual_include_questions = manual_export_format == "🎤 Interview Format (Questions & Answers)"
                
                if st.button("Create Biography from File", type="primary", use_container_width=True):
                    bio_text, all_stories, author_name, story_num, chapter_num, total_words = create_beautiful_biography(
                        uploaded_data, manual_include_questions
                    )
                    
                    safe_name = author_name.replace(" ", "_")
                    file_suffix = "_Interview" if manual_include_questions else "_Biography"
                    
                    # Create all formats
                    html_bio, _ = create_html_biography(uploaded_data, manual_include_questions)
                    md_bio = bio_text.replace("=" * 70, "#" * 3)
                    
                    # Try DOCX
                    docx_data = None
                    if DOCX_AVAILABLE:
                        try:
                            docx_bytes, docx_name, docx_chapters, docx_stories, docx_words = create_docx_biography(
                                uploaded_data, manual_include_questions
                            )
                            docx_data = docx_bytes
                        except Exception as e:
                            st.warning(f"Could not create DOCX: {str(e)}")
                    
                    # Show download buttons in columns
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        st.download_button(
                            label="📄 TXT",
                            data=bio_text,
                            file_name=f"{safe_name}{file_suffix}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                    
                    with col2:
                        st.download_button(
                            label="🌐 HTML",
                            data=html_bio,
                            file_name=f"{safe_name}{file_suffix}.html",
                            mime="text/html",
                            use_container_width=True
                        )
                    
                    with col3:
                        st.download_button(
                            label="📝 MD",
                            data=md_bio,
                            file_name=f"{safe_name}{file_suffix}.md",
                            mime="text/markdown",
                            use_container_width=True
                        )
                    
                    with col4:
                        if docx_data:
                            st.download_button(
                                label="📘 DOCX",
                                data=docx_data,
                                file_name=f"{safe_name}{file_suffix}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        else:
                            st.button("📘 DOCX", disabled=True, use_container_width=True)
                    
                    show_celebration()
                    st.success(f"Biography created for {author_name}!")
                    
            except Exception as e:
                st.error(f"❌ Error processing file: {str(e)}")

# ============================================================================
# FOOTER
# ============================================================================
st.markdown("---")
st.caption("✨ **Tell My Story Biography Publisher** • Create beautiful books from your life stories • 4 Export Formats • 2 Export Styles • Professional formatting • Celebration effects included")
